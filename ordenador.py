import os
from pdf2image import convert_from_path
from docx import Document
from io import BytesIO
from PIL import Image
from docx.shared import Inches

def set_cell_width(cell, width):
    cell.width = width
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.left_indent = Inches(0)

def pdfs_to_single_docx(folder_path, docx_path, crop_dimensions, images_per_row):
    doc = Document()
    # Genero una tabla para poder colocar uno al lado del otro
    table = doc.add_table(rows=1, cols=images_per_row)  # Inicia con una fila y la cantidad de columnas deseada
    table.autofit = False

    current_image = 0
    row_cells = table.rows[0].cells  # Referencia a las celdas de la primera fila

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.1)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.9)

    for filename in os.listdir(folder_path):
        if filename.endswith('.pdf'):
            pdf_path = os.path.join(folder_path, filename)
            images = convert_from_path(pdf_path)

            for image in images:
                if current_image == images_per_row:  # Si se alcanza el límite, crea una nueva fila
                    row_cells = table.add_row().cells
                    current_image = 0  # Reinicia el contador de imágenes

                image_cropped = image.crop(crop_dimensions)
                new_size = (int(image_cropped.width), int(image_cropped.height))
                image_resized = image_cropped.resize(new_size, Image.Resampling.LANCZOS)

                image_stream = BytesIO()
                image_resized.save(image_stream, format='PNG')
                image_stream.seek(0)

                # Insertar imagen en la celda correspondiente
                paragraph = row_cells[current_image].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(1.8), height=Inches(2.3)) # Ancho - Alto

                current_image += 1

    # Ajustar visualmente la tabla para que no muestre bordes
    width_in_inches = 1.9
    for row in table.rows:
        for cell in row.cells:
            set_cell_width(cell, Inches(width_in_inches))

    #if os.path.exists(docx_path):
    try:
        doc.save(docx_path)
        print("El documento se guardó correctamente.")
    except PermissionError:
        print("Error al guardar el documento. Cierre el documento que tiene abierto y vuelva a ejecutar.")
    except os.path.exists:
        print("Error al guardar el documento. Elimine el anterior documento y vuelva a intentar.")

# Configuración de las rutas y dimensiones
folder_path = r'..\Automatizacion_Comprobantes\Comprobantes'
docx_path = r'..\Automatizacion_Comprobantes\Resultados\documento.docx'
crop_dimensions = (150, 150, 850, 1200)  # Ajusta las dimensiones de recorte
images_per_row = 4  # Ajusta esto basado en el tamaño de tus imágenes y el ancho de la página

print("Creando archivo...")
pdfs_to_single_docx(folder_path, docx_path, crop_dimensions, images_per_row)
